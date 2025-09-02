import yaml
import re
from typing import Dict, List, Any, Optional
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

class ImmutaRuleExplainer:
    def __init__(self):
        self.rules = []
        self.explanations = []
    
    def parse_yaml_file(self, file_path: str) -> Dict[str, Any]:
        """Parse YAML configuration file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read().strip()
                if not content:
                    print(f"Warning: Empty YAML file {file_path}")
                    return {}
                
                # Fix common YAML formatting issues
                # Replace tabs with spaces
                content = content.replace('\t', '    ')
                
                # Fix Windows line endings
                content = content.replace('\r\n', '\n')
                
                return yaml.safe_load(content)
        except yaml.YAMLError as e:
            print(f"YAML parsing error in {file_path}: {e}")
            # Try to provide more specific error info
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    lines = file.readlines()
                    print(f"File has {len(lines)} lines")
                    print(f"First few lines: {lines[:3]}")
            except:
                pass
            return {}
        except Exception as e:
            print(f"Error reading file {file_path}: {e}")
            return {}
    
    def extract_rules(self, config: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Extract rules from configuration"""
        rules = []
        
        # Check for direct rules
        if 'rules' in config:
            rules.extend(config['rules'])
        
        # Check for rules within actions
        if 'actions' in config:
            for action in config['actions']:
                if 'rules' in action:
                    rules.extend(action['rules'])
        
        return rules
    
    def explain_predicate(self, predicate: str) -> str:
        """Convert predicate logic to human-readable explanation"""
        explanations = []
        
        # Handle split operations
        if 'split(' in predicate:
            pattern = r"split\(\s*([^,]+),\s*'([^']+)'\s*\)\s*\[safe_offset\((\d+)\)\]\s*in\s*\(\s*([^)]+)\s*\)"
            match = re.search(pattern, predicate)
            if match:
                field, delimiter, offset, values = match.groups()
                clean_values = [v.strip().strip("'\"") for v in values.split(',')]
                explanations.append(f"the first part of {field.strip()} (split by '{delimiter}') is one of: {', '.join(clean_values)}")
        
        # Handle direct field comparisons
        if ' in (' in predicate and 'split(' not in predicate:
            pattern = r"(\w+)\s+in\s+\(\s*([^)]+)\s*\)"
            match = re.search(pattern, predicate)
            if match:
                field, values = match.groups()
                clean_values = [v.strip().strip("'\"") for v in values.split(',')]
                explanations.append(f"{field} is one of: {', '.join(clean_values)}")
        
        # Handle attributeValuesContains
        if '@attributeValuesContains' in predicate:
            pattern = r"@attributeValuesContains\(\s*'([^']+)',\s*'([^']+)'\s*\)"
            match = re.search(pattern, predicate)
            if match:
                attr1, attr2 = match.groups()
                explanations.append(f"the user's {attr1} matches values in {attr2}")
        
        return ' or '.join(explanations) if explanations else predicate
    
    def explain_rule(self, rule: Dict[str, Any], rule_index: int) -> str:
        """Generate step-by-step explanation for a single rule"""
        explanation = f"\n**Rule {rule_index + 1}:**\n"
        
        config = rule.get('config', {})
        predicate = config.get('predicate', '')
        matches = config.get('matches', [])
        
        # Check for inclusions and exceptions at rule level or config level
        inclusions = rule.get('inclusions', config.get('inclusions', {}))
        exceptions = rule.get('exceptions', config.get('exceptions', {}))
        operator = rule.get('operator', config.get('operator', 'any'))
        rule_type = rule.get('type', config.get('type', 'Unknown'))
        

        
        # Handle inclusions
        if inclusions:
            explanation += "**Step 1: Check Inclusions**\n"
            
            # Check attributes
            attributes = inclusions.get('attributes', [])
            groups = inclusions.get('groups', [])
            
            conditions = []
            if attributes:
                for attr in attributes:
                    attr_name = attr.get('name', '')
                    attr_value = attr.get('value', '')
                    conditions.append(f"user's {attr_name} is '{attr_value}'")
            
            if groups:
                conditions.append(f"user belongs to one of these groups: {', '.join(groups)}")
            
            if conditions:
                if operator == 'any':
                    explanation += f"Immuta checks if {' OR '.join(conditions)}.\n"
                else:
                    explanation += f"Immuta checks if {' AND '.join(conditions)}.\n"
                
                predicate_explanation = self.explain_predicate(predicate)
                explanation += f"- **Action if True:** User will see data where {predicate_explanation}.\n"
                explanation += f"- **Action if False:** Move to next condition.\n\n"
        
        # Handle exceptions
        if exceptions:
            explanation += "**Step 2: Check Exceptions**\n"
            exception_groups = exceptions.get('groups', [])
            if exception_groups:
                explanation += f"Immuta checks if user belongs to exception groups: {', '.join(exception_groups)}.\n"
                explanation += f"- **Action if Yes:** User will see all data (exception applies).\n"
                explanation += f"- **Action if No:** Apply the standard rule filter.\n\n"
        
        # Handle User Entitlements rules with matches
        if matches and rule_type == 'Row Restriction by User Entitlements':
            explanation += "**User Entitlements Rule:**\n"
            for match in matches:
                attribute = match.get('attribute', '')
                tag = match.get('tag', '')
                match_type = match.get('type', '')
                explanation += f"User's {attribute} must match values in {tag} (type: {match_type}).\n"
            explanation += "\n"
        
        # Handle Masking rules
        elif rule_type == 'Masking':
            explanation += "**Masking Rule:**\n"
            fields = config.get('fields', [])
            masking_config = config.get('maskingConfig', {})
            masking_type = masking_config.get('type', 'Unknown')
            
            if fields:
                explanation += "This rule applies masking to the following fields:\n"
                for field in fields:
                    column_tag = field.get('columnTag', '')
                    field_type = field.get('type', '')
                    explanation += f"- {column_tag} (type: {field_type})\n"
                explanation += f"**Masking Type:** {masking_type}\n"
                explanation += f"**Action:** Data in these fields will be masked using {masking_type} method.\n\n"
        
        # If no inclusions, explain the predicate directly
        elif not inclusions and not exceptions and predicate:
            predicate_explanation = self.explain_predicate(predicate)
            explanation += f"**Condition:** User will see data where {predicate_explanation}.\n\n"
        
        # Handle rules with no specific conditions
        elif not inclusions and not exceptions and not predicate and not matches:
            explanation += "**Universal Rule:** This rule applies to all users and data.\n\n"
        
        return explanation
    
    def get_dataset_name(self, config: Dict) -> str:
        """Extract dataset name from YAML config"""
        # Try to get from circumstances tag
        circumstances = config.get('circumstances', [])
        for circ in circumstances:
            if circ.get('type') == 'tags':
                tag = circ.get('tag', '')
                if 'Table.' in tag:
                    # Extract table name from tag like "Data Entity.PO.Table.datamart_supplychain_procurement.bank_guarantee"
                    parts = tag.split('.')
                    if len(parts) >= 2:
                        return parts[-1]  # Get last part as table name
        
        # Fallback to policy name
        return config.get('name', 'unknown_dataset').replace(' ', '_').replace(':', '')
    
    def process_yaml_file(self, file_path: str) -> str:
        """Process a single YAML file and generate explanation"""
        config = self.parse_yaml_file(file_path)
        if not config:
            return f"# Error Processing File\n\nDataset/Table: Unknown\nFile Name: {os.path.basename(file_path)}\n\n## Error\n\nCould not parse YAML file. The file may be empty, corrupted, or contain invalid YAML syntax.\n\n## Troubleshooting\n\n- Check if the file is empty\n- Verify YAML syntax is correct\n- Ensure file encoding is UTF-8"
        
        rules = self.extract_rules(config)
        if not rules:
            dataset_name = self.get_dataset_name(config)
            return f"# Immuta Rule Configuration\n\nDataset/Table: {dataset_name}\nFile Name: {os.path.basename(file_path)}\n\n## Configuration\n\n```yaml\n{yaml.dump(config, default_flow_style=False, indent=2, sort_keys=False, allow_unicode=True)}```\n\n## Analysis\n\nNo rules found in this configuration file. This may be:\n- A configuration file without rules\n- A template or placeholder file\n- An incomplete configuration"
        
        dataset_name = self.get_dataset_name(config)
        explanation = f"# Immuta Rule Configuration Explanation\n"
        explanation += f"Dataset/Table: {dataset_name}\n"
        explanation += f"File Name: {os.path.basename(file_path)}\n\n"
        
        explanation += "## Configuration\n"
        explanation += "```yaml\n"
        explanation += yaml.dump(config, default_flow_style=False, indent=2, sort_keys=False, allow_unicode=True)
        explanation += "```\n\n"
        
        explanation += "## Explanation\n"
        
        for i, rule in enumerate(rules):
            explanation += self.explain_rule(rule, i)
        
        return explanation
    
    def generate_docx(self, content: str, output_path: str):
        """Generate Word document with YAML and explanation"""
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.shared import OxmlElement, qn
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
        
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Create professional styles with consistent spacing
        styles = doc.styles
        
        # Title style
        title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Calibri'
        title_font.size = Pt(20)
        title_font.bold = True
        title_font.color.rgb = RGBColor(0, 120, 212)
        title_style.paragraph_format.space_after = Pt(18)
        title_style.paragraph_format.space_before = Pt(0)
        
        # Section heading style
        section_style = styles.add_style('SectionHeading', WD_STYLE_TYPE.PARAGRAPH)
        section_font = section_style.font
        section_font.name = 'Calibri'
        section_font.size = Pt(16)
        section_font.bold = True
        section_font.color.rgb = RGBColor(0, 120, 212)
        section_style.paragraph_format.space_after = Pt(12)
        section_style.paragraph_format.space_before = Pt(18)
        
        # Rule heading style
        rule_style = styles.add_style('RuleHeading', WD_STYLE_TYPE.PARAGRAPH)
        rule_font = rule_style.font
        rule_font.name = 'Calibri'
        rule_font.size = Pt(14)
        rule_font.bold = True
        rule_font.color.rgb = RGBColor(68, 114, 196)
        rule_style.paragraph_format.space_after = Pt(6)
        rule_style.paragraph_format.space_before = Pt(12)
        
        # Step heading style
        step_style = styles.add_style('StepHeading', WD_STYLE_TYPE.PARAGRAPH)
        step_font = step_style.font
        step_font.name = 'Calibri'
        step_font.size = Pt(12)
        step_font.bold = True
        step_font.color.rgb = RGBColor(112, 173, 71)
        step_style.paragraph_format.space_after = Pt(6)
        step_style.paragraph_format.space_before = Pt(9)
        
        # Body text style
        body_style = styles.add_style('BodyText', WD_STYLE_TYPE.PARAGRAPH)
        body_font = body_style.font
        body_font.name = 'Calibri'
        body_font.size = Pt(11)
        body_style.paragraph_format.space_after = Pt(6)
        body_style.paragraph_format.line_spacing = 1.15
        
        # Bullet style
        bullet_style = styles.add_style('CustomBullet', WD_STYLE_TYPE.PARAGRAPH)
        bullet_font = bullet_style.font
        bullet_font.name = 'Calibri'
        bullet_font.size = Pt(11)
        bullet_style.paragraph_format.left_indent = Inches(0.25)
        bullet_style.paragraph_format.space_after = Pt(3)
        bullet_style.paragraph_format.line_spacing = 1.15
        
        # Add MFEC logo
        try:
            logo_path = os.path.join(os.path.dirname(__file__), 'LogoMFEC.png')
            if os.path.exists(logo_path):
                from PIL import Image as PILImage
                pil_img = PILImage.open(logo_path)
                img_width, img_height = pil_img.size
                aspect_ratio = img_width / img_height
                
                max_width = Inches(1.5)
                logo_height = max_width / aspect_ratio
                
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                run = paragraph.add_run()
                run.add_picture(logo_path, width=max_width)
                
                doc.add_paragraph()
        except:
            pass
        
        # Add title with modern styling
        title = doc.add_paragraph('Immuta Rule Configuration Analysis', style='CustomTitle')
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add spacing
        doc.add_paragraph()
        
        # Extract info from content
        lines = content.split('\n')
        dataset_name = "Unknown"
        file_name = "Unknown"
        
        for line in lines:
            if line.startswith('Dataset/Table:'):
                dataset_name = line.replace('Dataset/Table:', '').strip()
            elif line.startswith('File Name:'):
                file_name = line.replace('File Name:', '').strip()
        
        # Add document info section with professional table
        info_table = doc.add_table(rows=3, cols=2)
        info_table.style = 'Light Grid Accent 1'
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        header_cells = info_table.rows[0].cells
        header_cells[0].text = 'Document Information'
        header_cells[1].text = ''
        
        # Merge header cells
        header_cells[0].merge(header_cells[1])
        
        # Style header
        header_para = header_cells[0].paragraphs[0]
        header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        header_run = header_para.runs[0]
        header_run.font.name = 'Segoe UI'
        header_run.font.size = Pt(12)
        header_run.font.bold = True
        header_run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Set header background
        shading_elm = parse_xml(r'<w:shd {} w:fill="0078D4"/>'.format(nsdecls('w')))
        header_cells[0]._tc.get_or_add_tcPr().append(shading_elm)
        
        # Data rows
        info_table.cell(1, 0).text = 'Dataset/Table:'
        info_table.cell(1, 1).text = dataset_name
        info_table.cell(2, 0).text = 'File Name:'
        info_table.cell(2, 1).text = file_name
        
        # Style data rows
        for i in range(1, 3):
            for j in range(2):
                cell = info_table.cell(i, j)
                para = cell.paragraphs[0]
                run = para.runs[0] if para.runs else para.add_run(cell.text)
                run.font.name = 'Segoe UI'
                run.font.size = Pt(10)
                
                if j == 0:  # First column - labels
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 120, 212)
                    # Light blue background for labels
                    shading_elm = parse_xml(r'<w:shd {} w:fill="E7F3FF"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shading_elm)
        
        # Add spacing
        doc.add_paragraph()
        
        # Process content sections
        sections = content.split('\n## ')
        
        for section in sections:
            if section.startswith('Configuration'):
                # YAML Configuration Section
                config_heading = doc.add_paragraph('YAML Configuration', style='SectionHeading')
                doc.add_paragraph()  # Spacing
                
                yaml_start = section.find('```yaml')
                yaml_end = section.find('```', yaml_start + 7)
                if yaml_start != -1 and yaml_end != -1:
                    yaml_content = section[yaml_start + 7:yaml_end].strip()
                    
                    # Create professional YAML display table
                    yaml_table = doc.add_table(rows=1, cols=1)
                    yaml_table.style = 'Table Grid'
                    yaml_cell = yaml_table.cell(0, 0)
                    
                    # Clear default paragraph and add YAML content
                    yaml_cell.paragraphs[0].clear()
                    yaml_para = yaml_cell.paragraphs[0]
                    yaml_run = yaml_para.add_run(yaml_content)
                    yaml_run.font.name = 'Consolas'
                    yaml_run.font.size = Pt(9)
                    
                    # Set cell background and padding
                    shading_elm = parse_xml(r'<w:shd {} w:fill="F8F8F8"/>'.format(nsdecls('w')))
                    yaml_cell._tc.get_or_add_tcPr().append(shading_elm)
                    
                    # Add border styling
                    tc = yaml_cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcBorders = OxmlElement('w:tcBorders')
                    for border_name in ['top', 'left', 'bottom', 'right']:
                        border = OxmlElement(f'w:{border_name}')
                        border.set(qn('w:val'), 'single')
                        border.set(qn('w:sz'), '4')
                        border.set(qn('w:space'), '0')
                        border.set(qn('w:color'), 'CCCCCC')
                        tcBorders.append(border)
                    tcPr.append(tcBorders)
                    
            elif section.startswith('Explanation') or 'Rule' in section:
                # Rule Explanations Section
                explain_heading = doc.add_paragraph('Rule Explanations', style='SectionHeading')
                doc.add_paragraph()  # Spacing
                
                explanation_content = section[len('Explanation\n'):] if section.startswith('Explanation') else section
                
                lines = explanation_content.split('\n')
                current_rule_number = 0
                
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                        
                    if line.startswith('**Rule '):
                        current_rule_number += 1
                        # Add spacing before new rule (except first)
                        if current_rule_number > 1:
                            doc.add_paragraph()
                        
                        rule_text = line.strip('*')
                        rule_para = doc.add_paragraph(rule_text, style='RuleHeading')
                        
                        # Add rule number box
                        rule_table = doc.add_table(rows=1, cols=1)
                        rule_table.style = 'Table Grid'
                        rule_cell = rule_table.cell(0, 0)
                        rule_cell.text = f"Rule {current_rule_number}"
                        
                        # Style rule number box
                        rule_cell_para = rule_cell.paragraphs[0]
                        rule_cell_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        rule_cell_run = rule_cell_para.runs[0]
                        rule_cell_run.font.name = 'Segoe UI'
                        rule_cell_run.font.size = Pt(11)
                        rule_cell_run.font.bold = True
                        rule_cell_run.font.color.rgb = RGBColor(255, 255, 255)
                        
                        # Blue background for rule number
                        shading_elm = parse_xml(r'<w:shd {} w:fill="4472C4"/>'.format(nsdecls('w')))
                        rule_cell._tc.get_or_add_tcPr().append(shading_elm)
                        
                    elif line.startswith('**Step ') or line.startswith('**User ') or line.startswith('**Masking ') or line.startswith('**Condition:') or line.startswith('**Universal Rule:'):
                        step_text = line.strip('*')
                        step_para = doc.add_paragraph(step_text, style='StepHeading')
                        
                    elif line.startswith('- **'):
                        # Bullet points with enhanced formatting
                        p = doc.add_paragraph(style='List Bullet')
                        p.paragraph_format.left_indent = Inches(0.25)
                        
                        # Parse bold text within bullet points
                        text = line[2:]  # Remove '- '
                        parts = text.split('**')
                        
                        for i, part in enumerate(parts):
                            if part:
                                run = p.add_run(part)
                                run.font.name = 'Segoe UI'
                                run.font.size = Pt(10)
                                if i % 2 == 1:  # Odd indices are bold
                                    run.bold = True
                                    run.font.color.rgb = RGBColor(0, 120, 212)
                                    
                    elif line.startswith('- '):
                        # Regular bullet points
                        p = doc.add_paragraph(line[2:], style='List Bullet')
                        p.paragraph_format.left_indent = Inches(0.25)
                        run = p.runs[0]
                        run.font.name = 'Segoe UI'
                        run.font.size = Pt(10)
                        
                    elif not line.startswith('#') and line:
                        # Regular text with proper formatting
                        p = doc.add_paragraph(line, style='BodyText')
                        
                        # Add indentation for action descriptions
                        if 'Action if' in line or 'Immuta checks' in line:
                            p.paragraph_format.left_indent = Inches(0.25)
                            p.paragraph_format.space_after = Pt(6)
        
        # Add footer with generation info
        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        footer_run = footer_para.add_run('Generated by Immuta Rule Configuration Explainer')
        footer_run.font.name = 'Segoe UI'
        footer_run.font.size = Pt(8)
        footer_run.font.italic = True
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        
        doc.save(output_path)
        print(f"Professional document saved to: {output_path}")
    
    def generate_pdf(self, content: str, output_path: str):
        """Generate PDF document using reportlab"""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib.colors import HexColor
            
            doc = SimpleDocTemplate(output_path, pagesize=letter, 
                                  topMargin=1*inch, bottomMargin=1*inch, 
                                  leftMargin=1*inch, rightMargin=1*inch)
            styles = getSampleStyleSheet()
            story = []
            
            # Add MFEC logo if exists
            try:
                logo_path = os.path.join(os.path.dirname(__file__), 'LogoMFEC.png')
                if os.path.exists(logo_path):
                    # Get image dimensions and calculate proper scaling
                    from PIL import Image as PILImage
                    pil_img = PILImage.open(logo_path)
                    img_width, img_height = pil_img.size
                    aspect_ratio = img_width / img_height
                    
                    # Set smaller width and calculate height maintaining aspect ratio
                    max_width = 1*inch
                    logo_height = max_width / aspect_ratio
                    
                    # Create right-aligned logo
                    logo = Image(logo_path, width=max_width, height=logo_height)
                    logo.hAlign = 'RIGHT'
                    story.append(logo)
                    story.append(Spacer(1, 0.2*inch))
            except:
                pass
            
            # Custom styles for professional layout
            title_style = ParagraphStyle('CustomTitle', parent=styles['Title'], 
                                       fontSize=16, textColor=HexColor('#2C3E50'),
                                       spaceAfter=0.3*inch, alignment=1)
            
            heading1_style = ParagraphStyle('CustomHeading1', parent=styles['Heading1'],
                                          fontSize=14, textColor=HexColor('#34495E'),
                                          spaceAfter=0.2*inch, spaceBefore=0.3*inch)
            
            normal_style = ParagraphStyle('CustomNormal', parent=styles['Normal'],
                                        fontSize=11, leading=14, textColor=HexColor('#2C3E50'))
            
            # Create info table
            from reportlab.platypus import Table, TableStyle
            from reportlab.lib import colors
            
            # Extract info from content
            dataset_name = "Unknown"
            file_name = "Unknown"
            lines = content.split('\n')
            
            for line in lines:
                if line.startswith('Dataset/Table:'):
                    dataset_name = line.replace('Dataset/Table:', '').strip()
                elif line.startswith('File Name:'):
                    file_name = line.replace('File Name:', '').strip()
            
            # Add info table
            info_data = [['Dataset/Table:', dataset_name], ['File Name:', file_name]]
            info_table = Table(info_data, colWidths=[2*inch, 4*inch])
            info_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), HexColor('#E7F3FF')),
                ('TEXTCOLOR', (0, 0), (0, -1), HexColor('#0078d4')),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('GRID', (0, 0), (-1, -1), 1, colors.lightgrey),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('LEFTPADDING', (0, 0), (-1, -1), 8),
                ('RIGHTPADDING', (0, 0), (-1, -1), 8),
            ]))
            story.append(info_table)
            story.append(Spacer(1, 0.3*inch))
            

            
            # Convert content to PDF with better formatting
            in_yaml_block = False
            for line in lines:
                if line.startswith('# '):
                    story.append(Paragraph(line[2:], title_style))
                elif line.startswith('## '):
                    story.append(Paragraph(line[3:], heading1_style))
                elif line.startswith('```yaml'):
                    in_yaml_block = True
                    continue
                elif line.startswith('```') and in_yaml_block:
                    in_yaml_block = False
                    story.append(Spacer(1, 0.2*inch))
                    continue
                elif in_yaml_block:
                    # YAML content with proper line spacing to prevent overlap
                    yaml_style = ParagraphStyle('YAMLStyle', parent=styles['Normal'],
                                               fontName='Courier', fontSize=8,
                                               leftIndent=15, backColor=HexColor('#F8F8F8'),
                                               borderWidth=0, borderColor=colors.lightgrey,
                                               borderPadding=12, leading=16,
                                               spaceBefore=4, spaceAfter=4)
                    # Preserve original spacing and indentation with better formatting
                    formatted_line = line.replace('    ', '&nbsp;&nbsp;&nbsp;&nbsp;').replace('  ', '&nbsp;&nbsp;')
                    # Add extra spacing for better readability
                    if formatted_line.strip():
                        story.append(Paragraph(formatted_line, yaml_style))
                        story.append(Spacer(1, 0.02*inch))  # Small spacer between lines
                elif line.startswith('**Rule '):
                    rule_style = ParagraphStyle('RuleStyle', parent=styles['Heading2'], 
                                              fontSize=12, textColor=HexColor('#4472C4'), 
                                              spaceAfter=0.1*inch, spaceBefore=0.15*inch,
                                              fontName='Helvetica-Bold')
                    story.append(Paragraph(line[2:-2], rule_style))
                elif line.startswith('**Step ') or line.startswith('**User ') or line.startswith('**Masking ') or line.startswith('**Condition:') or line.startswith('**Universal Rule:'):
                    step_style = ParagraphStyle('StepStyle', parent=styles['Heading3'], 
                                              fontSize=11, textColor=HexColor('#70AD47'),
                                              spaceAfter=0.08*inch, spaceBefore=0.1*inch,
                                              fontName='Helvetica-Bold')
                    story.append(Paragraph(line[2:-2], step_style))
                elif line.startswith('- **Action'):
                    action_style = ParagraphStyle('ActionStyle', parent=styles['Normal'],
                                                fontSize=10, leading=12, textColor=HexColor('#2C3E50'),
                                                leftIndent=20, spaceAfter=0.03*inch)
                    bullet_text = line[2:].replace('**', '<b>', 1).replace('**', '</b>', 1)
                    story.append(Paragraph(f'• {bullet_text}', action_style))
                elif line.startswith('- '):
                    action_style = ParagraphStyle('ActionStyle', parent=styles['Normal'],
                                                fontSize=10, leading=12, textColor=HexColor('#2C3E50'),
                                                leftIndent=20, spaceAfter=0.03*inch)
                    story.append(Paragraph(f'• {line[2:]}', action_style))
                elif 'Immuta checks' in line or 'Action if' in line:
                    action_style = ParagraphStyle('ActionStyle', parent=styles['Normal'],
                                                fontSize=10, leading=12, textColor=HexColor('#2C3E50'),
                                                leftIndent=20, spaceAfter=0.03*inch)
                    story.append(Paragraph(line, action_style))

                elif line.strip() and not line.startswith(('Dataset/Table:', 'File Name:')):
                    story.append(Paragraph(line, normal_style))
            
            doc.build(story)
            print(f"PDF saved to: {output_path}")
        except ImportError:
            raise Exception("reportlab not installed")
        except Exception as e:
            raise Exception(f"PDF generation failed: {e}")
    
    def _generate_pdf_reportlab(self, content: str, output_path: str):
        """Generate PDF using reportlab as fallback"""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib.units import inch
            
            doc = SimpleDocTemplate(output_path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Convert markdown-like content to PDF
            lines = content.split('\n')
            for line in lines:
                if line.startswith('# '):
                    story.append(Paragraph(line[2:], styles['Title']))
                elif line.startswith('## '):
                    story.append(Paragraph(line[3:], styles['Heading1']))
                elif line.startswith('**') and line.endswith('**'):
                    story.append(Paragraph(line[2:-2], styles['Heading2']))
                elif line.strip():
                    story.append(Paragraph(line, styles['Normal']))
                story.append(Spacer(1, 0.1*inch))
            
            doc.build(story)
            print(f"PDF saved to: {output_path}")
        except ImportError:
            raise Exception("reportlab not installed. Install with: pip install reportlab")
        except Exception as e:
            raise Exception(f"PDF generation failed: {e}")

def main():
    explainer = ImmutaRuleExplainer()
    
    # Get current directory
    current_dir = os.getcwd()
    
    # Find all YAML files
    yaml_files = [f for f in os.listdir(current_dir) if f.endswith('.yaml')]
    
    if not yaml_files:
        print("No YAML files found in current directory")
        return
    
    print("Available YAML files:")
    for i, file in enumerate(yaml_files, 1):
        print(f"{i}. {file}")
    
    try:
        choice = input("\nEnter file number to process (or 'all' for all files): ").strip()
        
        if choice.lower() == 'all':
            for yaml_file in yaml_files:
                print(f"\nProcessing {yaml_file}...")
                explanation = explainer.process_yaml_file(yaml_file)
                
                # Generate DOCX
                output_file = yaml_file.replace('.yaml', '_explanation.docx')
                explainer.generate_docx(explanation, output_file)
        else:
            file_index = int(choice) - 1
            if 0 <= file_index < len(yaml_files):
                selected_file = yaml_files[file_index]
                print(f"\nProcessing {selected_file}...")
                
                explanation = explainer.process_yaml_file(selected_file)
                print(explanation)
                
                # Generate DOCX
                output_file = selected_file.replace('.yaml', '_explanation.docx')
                explainer.generate_docx(explanation, output_file)
            else:
                print("Invalid file number")
    
    except ValueError:
        print("Invalid input")
    except KeyboardInterrupt:
        print("\nOperation cancelled")

if __name__ == "__main__":
    main()